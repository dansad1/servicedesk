from django.shortcuts import render, redirect, get_object_or_404
from django.conf import settings
from docx.document import Document

from service.forms.Chat_forms import DocumentUploadForm
from service.models import Doc    # Alias to differentiate from python-docx Document
import os
from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from django.shortcuts import render, get_object_or_404, redirect
from django.conf import settings

def upload_document(request):
    if request.method == 'POST':
        form = DocumentUploadForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('document_list')  # Перенаправление на список документов
    else:
        form = DocumentUploadForm()
    return render(request, 'documents/document_upload.html', {'form': form})
def document_list(request):
    documents = Doc.objects.all()
    return render(request, 'documents/document_list.html', {'documents': documents})
from docx import Document

def extract_data(doc_file_path):
    doc = Document(doc_file_path)
    text_content = [para.text for para in doc.paragraphs]
    tables = [[cell.text for cell in row.cells] for table in doc.tables for row in table.rows]
    images = []
    for rel in doc.part.rels.values():
        if 'image' in rel.reltype:
            images.append(rel.target_part.blob)
    return text_content, tables, images


def analyze_document(request, doc_id):
    document = get_object_or_404(Doc, id=doc_id)
    if not document.analyzed:
        text_content, tables, images = extract_data(document.doc_file.path)
        document.text_content = '\n'.join(text_content)  # Saving extracted text
        document.image_urls = '\n'.join([str(img) for img in images])  # Simplified for example
        document.tables = '\n'.join(['; '.join(row) for table in tables for row in table])  # Simplified for example
        document.analyzed = True
        document.save()
    else:
        text_content = document.text_content.split('\n')
        tables = [row.split(';') for row in document.tables.split('\n')]
        images = document.image_urls.split('\n')

    context = {
        'document': document,
        'text_content': text_content,
        'tables': tables,
        'images': images
    }
    return render(request, 'documents/document_analysis.html', context)
